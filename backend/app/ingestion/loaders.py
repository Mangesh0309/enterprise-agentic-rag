from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass(slots=True)
class ParsedPage:
    text: str
    page_number: int | None = None
    metadata: dict | None = None


@dataclass(slots=True)
class ParsedDocument:
    pages: list[ParsedPage]
    metadata: dict


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())


def _read_with_llama_index(path: Path) -> list[ParsedPage]:
    from llama_index.core import SimpleDirectoryReader

    documents = SimpleDirectoryReader(input_files=[str(path)]).load_data()
    pages: list[ParsedPage] = []
    for index, doc in enumerate(documents, start=1):
        metadata = dict(doc.metadata or {})
        page_label = metadata.get("page_label") or metadata.get("page_number")
        try:
            page_number = int(page_label) if page_label is not None else index
        except (TypeError, ValueError):
            page_number = index
        pages.append(ParsedPage(text=doc.text, page_number=page_number, metadata=metadata))
    return pages


def parse_document(path: str, content_type: str) -> ParsedDocument:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    metadata = {"filename": file_path.name, "content_type": content_type}

    if suffix == ".docx":
        text = _read_docx(file_path)
        return ParsedDocument(pages=[ParsedPage(text=text, page_number=1)], metadata=metadata)

    if suffix in {".txt", ".md", ".csv"}:
        text = _read_text(file_path)
        return ParsedDocument(pages=[ParsedPage(text=text, page_number=1)], metadata=metadata)

    pages = _read_with_llama_index(file_path)
    return ParsedDocument(pages=pages, metadata=metadata)
